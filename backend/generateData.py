import os
import subprocess
import pandas as pd
from paddleocr import PaddleOCR
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import Docx2txtLoader, TextLoader, PyPDFLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# =======================
# DOCX表格处理器
# =======================
def process_docx_tables_as_documents(docx_path, doc_type="Unknown"):
    """
    处理DOCX文件：表格单独成块，文本按800字符分块

    策略：
    1. 按顺序读取文档元素
    2. 文本内容累积到800字符时立即分块保存
    3. 遇到表格时，立即保存当前累积的文本块（无论多少字符）
    4. 表格单独作为完整块，前面添加文件名

    Args:
        docx_path: DOCX文件路径
        doc_type: 文档类型

    Returns:
        文档列表（表格块+文本分块）
    """
    documents = []

    try:
        doc = DocxDocument(docx_path)

        current_text_chunk = ""
        table_count = 0
        text_chunk_count = 0
        last_paragraph_text = ""  # 记录最后一段的文本，作为表格名称

        element_idx = 0

        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = element
                text = para.text.strip()
                if text:
                    last_paragraph_text = text  # 更新最后段落文本
                    current_text_chunk += text + "\n"

                    # 立即检查是否超过800字符，超过则分块
                    while len(current_text_chunk) >= 800:
                        text_chunk_count += 1
                        chunk_content = current_text_chunk[:800]
                        chunk_content = add_filename_prefix(chunk_content, docx_path, doc_type)

                        documents.append(Document(
                            page_content=chunk_content,
                            metadata={
                                "source": docx_path,
                                "type": doc_type,
                                "chunk_index": text_chunk_count,
                                "content_type": "text_content",
                                "length": len(chunk_content)
                            }
                        ))
                        current_text_chunk = current_text_chunk[800:]

            elif isinstance(element, CT_Tbl):
                table = doc.tables[element_idx]
                element_idx += 1

                # 遇到表格时，立即保存当前累积的文本块（无论多少字符）
                if current_text_chunk:
                    text_chunk_count += 1
                    chunk_content = current_text_chunk
                    chunk_content = add_filename_prefix(chunk_content, docx_path, doc_type)

                    documents.append(Document(
                        page_content=chunk_content,
                        metadata={
                            "source": docx_path,
                            "type": doc_type,
                            "chunk_index": text_chunk_count,
                            "content_type": "text_content",
                            "length": len(chunk_content)
                        }
                    ))
                    current_text_chunk = ""

                # 处理表格作为单独的完整块
                table_count += 1
                table_content = extract_table_content(table)
                filename = os.path.basename(docx_path)
                table_name = last_paragraph_text.strip() if last_paragraph_text.strip() else f"表格{table_count}"
                full_table_content = f"[{filename}]\n[{table_name}]\n{table_content}"

                metadata = {
                    "source": docx_path,
                    "type": doc_type,
                    "table_index": table_count,
                    "table_name": last_paragraph_text,
                    "content_type": "structured_table",
                    "rows": len(table.rows),
                    "cols": len(table.columns) if hasattr(table, 'columns') and table.columns else len(table.rows[0].cells) if table.rows else 0,
                    "has_context": False
                }

                documents.append(Document(
                    page_content=full_table_content,
                    metadata=metadata
                ))

        # 处理最后的文本块
        if current_text_chunk:
            text_chunk_count += 1
            chunk_content = current_text_chunk
            chunk_content = add_filename_prefix(chunk_content, docx_path, doc_type)

            documents.append(Document(
                page_content=chunk_content,
                metadata={
                    "source": docx_path,
                    "type": doc_type,
                    "chunk_index": text_chunk_count,
                    "content_type": "text_content",
                    "length": len(chunk_content)
                }
            ))

    except Exception as e:
        print(f"DOCX表格处理失败: {docx_path}, 原因: {e}")

    return documents



def add_filename_prefix(content, file_path, doc_type):
    """
    为技术文档的分块添加文件名前缀

    Args:
        content: 原始内容
        file_path: 文件路径
        doc_type: 文档类型

    Returns:
        添加前缀后的内容
    """
    if doc_type.lower() == "technology":
        filename = os.path.basename(file_path)
        return f"[{filename}]\n{content}"
    return content

def extract_table_content(table):
    """
    提取表格内容为字符串格式

    Args:
        table: docx表格对象

    Returns:
        表格内容的字符串表示
    """
    content_lines = []

    # 添加表格数据（不单独处理表头，因为规格书中的表格可能没有明确的表头）
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            # 处理单元格内容，保留换行符
            cell_text = cell.text.strip()
            if not cell_text:
                cell_text = "[空]"
            row_data.append(cell_text)

        content_lines.append(f"行{row_idx + 1}: {' | '.join(row_data)}")

    # 返回完整的表格内容
    return "\n".join(content_lines)

# =======================
# OCR 初始化
# =======================
ocr = PaddleOCR(use_angle_cls=True, lang='ch', use_gpu=True)

def ocr_image_to_txt(img_path):
    """对图片执行 OCR 并生成同名 TXT 文件"""
    txt_file = os.path.splitext(img_path)[0] + ".txt"
    if os.path.exists(txt_file):
        return txt_file

    result = ocr.ocr(img_path, cls=True)
    with open(txt_file, "w", encoding="utf-8") as f:
        if result:
            for line in result[0]:
                f.write(line[1][0] + "\n")
    return txt_file

def excel_to_txt(file_path):
    """将 Excel 文件转换为 TXT，每行拼接单元格"""
    try:
        xls = pd.ExcelFile(file_path)
        texts = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            for row in df.itertuples(index=False):
                row_text = "\t".join([str(cell) for cell in row if pd.notna(cell)])
                if row_text.strip():
                    texts.append(row_text)
        txt_file = os.path.splitext(file_path)[0] + ".txt"
        with open(txt_file, "w", encoding="utf-8") as f:
            f.write("\n".join(texts))
        return txt_file
    except Exception as e:
        print(f"Excel 处理失败: {file_path}, 原因: {e}")
        return None

def csv_to_txt(file_path):
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin1']
    
    # 检测编码
    try:
        import chardet
        with open(file_path, 'rb') as f:
            encodings.insert(0, chardet.detect(f.read())['encoding'])
    except: pass
    
    # 尝试读取
    for enc in encodings:
        try:
            # 检测标题
            with open(file_path, 'r', encoding=enc, errors='replace') as f:
                has_header = ',' in f.readline()
            
            # 读取转换
            df = pd.read_csv(file_path, header=(0 if has_header else None), encoding=enc, on_bad_lines='skip')
            texts = []
            
            # 处理数据
            if has_header and len(df.columns) > 0:
                texts.append("\t".join(str(col) for col in df.columns))
            
            for row in df.itertuples(index=False):
                text = "\t".join(str(cell) for cell in row if pd.notna(cell))
                if text: texts.append(text)
            
            # 保存
            txt_file = os.path.splitext(file_path)[0] + ".txt"
            with open(txt_file, "w", encoding="utf-8") as f:
                f.write("\n".join(texts))
            return txt_file
        except Exception:
            continue
    
    return None


text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,          # 增加到 800 字符，保证每个块有完整的信息
    chunk_overlap=150,        # 增加重叠，避免关键信息在边界丢失
    length_function=len,
    separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", "；", ";", "，", ","]  # 优先按段落和句子分割
)

os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
embedding_model = HuggingFaceEmbeddings(
    model_name="/New_Disk/liziwei/maidalun1020/bce-embedding-base_v1",
    model_kwargs={"device": "cuda"},
    encode_kwargs={"normalize_embeddings": True}
)

dir_path = '/home/liziwei/Emergency-LLM/backend/resource'
all_documents = []

for foldername, subfolders, filenames in os.walk(dir_path):
    for filename in filenames:
        file_path = os.path.join(foldername, filename)
        loader = None
        lower_name = filename.lower()

        if lower_name.endswith((".png", ".jpg", ".jpeg", ".bmp")):
            txt_file = ocr_image_to_txt(file_path)
            loader = TextLoader(txt_file, encoding='utf-8')
        elif lower_name.endswith((".xls", ".xlsx")):
            txt_file = excel_to_txt(file_path)
            loader = TextLoader(txt_file, encoding='utf-8') if txt_file else None
        elif lower_name.endswith(".csv"):
            txt_file = csv_to_txt(file_path)
            if txt_file:
                loader = TextLoader(txt_file, encoding='utf-8')
            else:
                continue
        elif lower_name.endswith(".pdf"):
            print("加载 PDF:", file_path)
            loader = PyPDFLoader(file_path)
        elif lower_name.endswith(".docx"):
            print("处理 DOCX 表格:", file_path)
            # 特殊处理：将表格作为整体文档处理
            docs = process_docx_tables_as_documents(file_path, doc_type)
            # 直接添加到all_documents，不使用loader
            if docs:
                all_documents.extend(docs)
            continue  # 跳过下面的loader处理
        elif lower_name.endswith(".txt"):
            # 尝试多种编码加载 TXT 文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin1']
            try:
                import chardet
                with open(file_path, 'rb') as f:
                    detected = chardet.detect(f.read())
                    if detected['encoding']:
                        encodings.insert(0, detected['encoding'])
            except: pass
            
            for enc in encodings:
                try:
                    loader = TextLoader(file_path, encoding=enc)
                    # 尝试加载验证编码是否正确
                    _ = loader.load()
                    break
                except:
                    continue
            else:
                print(f"无法识别文件编码，跳过: {file_path}")
                continue
        elif lower_name.endswith((".wps", ".doc")):
            out_file = os.path.splitext(file_path)[0] + ".docx"
            try:
                # 步骤1：转换文件
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", foldername, file_path],
                    check=True
                )

                # 步骤2：处理转换后的文件
                docs = process_docx_tables_as_documents(out_file, doc_type)

                # 步骤3：只有在成功处理的情况下才添加到文档列表并删除原始文件
                if docs:
                    all_documents.extend(docs)
                    print(f"✅ 成功处理文件: {filename} -> {len(docs)} 个文档")

                    # 可选：删除原始DOC文件（现在默认保留）
                    # os.remove(file_path)
                else:
                    print(f"⚠️  文件处理完成但未生成文档: {filename}")

                # 保留转换后的DOCX文件作为备份
                continue  # 跳过下面的loader处理

            except Exception as e:
                print(f"❌ 文件处理失败，保留原始文件: {file_path}, 原因: {e}")
                # 如果转换失败，尝试清理可能产生的临时文件
                if os.path.exists(out_file):
                    try:
                        os.remove(out_file)
                        print(f"清理临时文件: {out_file}")
                    except:
                        pass
                continue
        else:
            print("未知格式文件，跳过:", file_path)
            continue
        if loader:
            try:
                # 提取文档类型（子文件夹名称）
                rel_path = os.path.relpath(foldername, dir_path)
                doc_type = os.path.normpath(rel_path).split(os.sep)[0] if rel_path != '.' else 'Unknown'
                
                docs = loader.load_and_split(text_splitter)

                # 清洗和过滤文档
                for doc in docs:
                    content = doc.page_content.strip()

                    # 过滤条件
                    # 1. 长度太短（<20字符）
                    if len(content) < 20:
                        continue

                    # 2. 只过滤包含乱码替换字符的文档
                    # � (U+FFFD) 是 Unicode 替换字符，表示无法解码的字节
                    if '�' in content:
                        # 统计乱码字符的比例
                        garbled_count = content.count('�')
                        # 如果乱码字符超过5个，或者占比超过5%，则跳过
                        if garbled_count > 5 or (garbled_count / len(content)) > 0.05:
                            print(f"跳过乱码文档片段（包含{garbled_count}个�字符）: {content[:80]}...")
                            continue

                    # 3. 只包含标点符号和空格（没有实际内容）
                    has_content = any(c.isalnum() or '\u4e00' <= c <= '\u9fff' for c in content)
                    if not has_content:
                        continue

                    # 为技术文档添加文件名前缀
                    content = add_filename_prefix(content, file_path, doc_type)

                    doc.page_content = content
                    doc.metadata["source"] = file_path
                    doc.metadata["type"] = doc_type
                    all_documents.append(doc)
                    
            except Exception as e:
                print(f"加载失败: {file_path}, 原因: {e}")

if all_documents:
    vdb = Chroma.from_documents(
        documents=all_documents,
        embedding=embedding_model,
        persist_directory="/home/liziwei/Emergency-LLM/new_main/model/vdb"
    )
    # vdb.persist()
    print("向量数据库创建成功！")
else:
    print("没有找到任何可处理的文件。")