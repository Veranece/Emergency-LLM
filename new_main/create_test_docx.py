#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建测试DOCX文件，包含表格用于测试表格识别功能
"""

from docx import Document
from docx.shared import Inches

def create_test_docx_with_table():
    """创建一个包含表格的测试DOCX文件"""

    # 创建新文档
    doc = Document()

    # 添加标题
    doc.add_heading('应急管理表格示例', 0)

    # 添加段落
    doc.add_paragraph('以下是应急管理相关的表格数据：')

    # 创建表格
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '应急类型'
    hdr_cells[1].text = '响应等级'
    hdr_cells[2].text = '所需装备'
    hdr_cells[3].text = '注意事项'

    # 添加数据行
    data = [
        ['火灾', '一级', '消防车、灭火器、防护服', '立即疏散人员，切断电源'],
        ['洪水', '二级', '救生艇、防水装备、通讯设备', '转移至高处，避免涉水'],
        ['地震', '一级', '救援队、医疗设备、帐篷', '寻找安全避难所，保护头部'],
        ['台风', '二级', '加固材料、发电机、应急食品', '加固房屋，储备物资']
    ]

    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data

    # 保存文档
    doc.save('/home/liziwei/Emergency-LLM/new_main/resources/files/test_emergency_table.docx')
    print("测试DOCX文件已创建：test_emergency_table.docx")

if __name__ == '__main__':
    create_test_docx_with_table()
